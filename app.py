import streamlit as st
from langchain_openai import ChatOpenAI
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import json
import re
import io

st.set_page_config(page_title="ATS Resume Architect", page_icon="📄", layout="wide")

# --- 1. HELPERS: TEXT EXTRACTION ---
def extract_text(uploaded_file):
    if uploaded_file is None:
        return ""
    file_ext = uploaded_file.name.split('.')[-1].lower()
    try:
        if file_ext == 'pdf':
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "\n".join([page.get_text() for page in doc])
        elif file_ext == 'docx':
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception:
        return ""
    return ""

# --- 2. HELPERS: WORD (.DOCX) GENERATION ---
def create_docx(data):
    doc = Document()
    
    # Name and Header
    header = doc.add_paragraph()
    run = header.add_run(data.get("name", "RESUME").upper())
    run.bold = True
    run.font.size = Pt(16)
    header.alignment = 1 # Center
    
    contact = doc.add_paragraph()
    contact.add_run(f"{data.get('email', '')}  |  {data.get('phone', '')}  |  {data.get('location', '')}")
    contact.alignment = 1
    
    # Helper to add sections
    def add_section(title, content):
        if content:
            doc.add_heading(title.upper(), level=1)
            p = doc.add_paragraph(content)
            p.style.font.size = Pt(10)

    add_section("Professional Summary", data.get("Professional_Summary"))
    add_section("Core Competencies", data.get("Core_Competencies"))
    add_section("Work Experience", data.get("Work_Experience"))
    add_section("Education", data.get("Education"))
    add_section("Certifications", data.get("Certifications"))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

# --- 3. HELPERS: PDF GENERATION ---
def create_pdf(data):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, data.get("name", "RESUME").upper(), ln=True, align='C')
    
    # Contact
    pdf.set_font("Helvetica", '', 10)
    contact_info = f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}"
    pdf.cell(0, 5, contact_info, ln=True, align='C')
    pdf.ln(5)

    sections = [
        ("Professional Summary", "Professional_Summary"),
        ("Core Competencies", "Core_Competencies"),
        ("Work Experience", "Work_Experience"),
        ("Education", "Education"),
        ("Certifications", "Certifications")
    ]
    
    for title, key in sections:
        content = data.get(key, "")
        if content:
            pdf.set_font("Helvetica", 'B', 12)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 8, title.upper(), ln=True, fill=True)
            pdf.ln(2)
            pdf.set_font("Helvetica", '', 10)
            # Standardizing text for PDF safety
            clean_text = content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(4)
            
    return pdf.output(dest='S')

# --- 4. MAIN UI ---
st.title("🎯 ATS Friendly Resume Generator")
st.info("Paste the job description and your current resume. The AI will generate an optimized, high-score resume in both Word and PDF formats.")

with st.sidebar:
    st.header("🔑 API Setup")
    api_key = st.secrets.get("OPENAI_API_KEY") or st.text_input("OpenAI API Key", type="password")
    st.divider()
    st.markdown("### ATS Guidelines:")
    st.caption("1. No Columns/Tables\n2. Standard Headings\n3. Keyword Matching\n4. Simple Fonts")

# Inputs
c1, c2 = st.columns(2)
with c1:
    st.subheader("📋 Job Description")
    jd_upload = st.file_uploader("Upload JD", type=["pdf", "docx"], key="jd_up")
    jd_text = st.text_area("Or Paste JD here:", height=250, key="jd_txt")

with c2:
    st.subheader("📄 Current Resume")
    res_upload = st.file_uploader("Upload Resume", type=["pdf", "docx"], key="res_up")
    res_text = st.text_area("Or Paste Resume here:", height=250, key="res_txt")

if st.button("🚀 Optimize & Generate Files", type="primary"):
    # Content collection
    final_jd = jd_text if jd_text else extract_text(jd_upload)
    final_res = res_text if res_text else extract_text(res_upload)
    
    if not final_jd or not final_res or not api_key:
        st.error("Missing Job Description, Resume, or API Key.")
    else:
        try:
            with st.spinner("Analyzing keywords and restructuring for ATS..."):
                llm = ChatOpenAI(model="gpt-4o-mini", api_key=api_key, temperature=0.2)
                
                prompt = f"""
                You are a Senior Technical Recruiter. Optimize this resume for the provided Job Description (JD).
                
                RULES:
                1. Identify the most important keywords in the JD.
                2. Use those keywords naturally in the Summary, Competencies, and Experience.
                3. Rewrite Experience bullet points using the STAR method (Situation, Task, Action, Result).
                4. Output MUST be in JSON format.
                
                JD: {final_jd}
                RESUME: {final_res}
                
                FORMAT:
                {{
                  "name": "Full Name",
                  "email": "email@example.com",
                  "phone": "Phone Number",
                  "location": "City, State",
                  "Professional_Summary": "ATS optimized summary...",
                  "Core_Competencies": "Skill 1, Skill 2, Skill 3...",
                  "Work_Experience": "- Company A: Bullet points...\\n- Company B: Bullet points...",
                  "Education": "Degree, University",
                  "Certifications": "Relevant certs only"
                }}
                """
                
                response = llm.invoke(prompt)
                
                # Robust JSON cleaning
                json_str = re.search(r"\{.*\}", response.content, re.DOTALL).group()
                resume_data = json.loads(json_str)
                
                # Create the Files
                docx_data = create_docx(resume_data)
                pdf_data = create_pdf(resume_data)
                
                st.success("✅ Resume Optimized Successfully!")
                
                # Display Analysis Results
                with st.expander("🔍 View Optimized Content Preview"):
                    st.write(f"**Targeting Name:** {resume_data['name']}")
                    st.write(f"**Keywords Added:** {resume_data['Core_Competencies']}")
                
                # Download Buttons
                btn_c1, btn_c2 = st.columns(2)
                with btn_c1:
                    st.download_button(
                        label="📥 Download Word (.docx)",
                        data=docx_data,
                        file_name="ATS_Resume_Optimized.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with btn_c2:
                    st.download_button(
                        label="📥 Download PDF (.pdf)",
                        data=pdf_data,
                        file_name="ATS_Resume_Optimized.pdf",
                        mime="application/pdf"
                    )
                    
        except Exception as e:
            st.error(f"⚠️ An error occurred: {str(e)}")
